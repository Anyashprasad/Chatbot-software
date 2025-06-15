#dataset
import os
import pickle
import fitz
import re
import numpy as np
from flask import jsonify
from rapidfuzz import process, fuzz
import ollama
import pytesseract
from pdf2image import convert_from_path
import hashlib
import json
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
import threading


OCR_CACHE_DIR = "data/ocr_cache"
GLOSSARY_CACHE_DIR = "data/glossary_cache"
_glossary_cache = {}
_glossary_lock = threading.Lock()

def preprocess_query(query):
    stopwords = [
        "explain", "meaning of", "describe", "what is", "tell me about",
        "give details of", "give me details of"
    ]
    q = query.lower()
    for w in stopwords:
        q = q.replace(w, "")
    return q.strip()

def extract_dynamic_glossary(full_text):
    glossary = {}
    for match in re.finditer(r'([A-Za-z ,&/-]+)\s*\(\s*([A-Z]{2,})\s*\)', full_text):
        definition, acronym = match.group(1).strip(), match.group(2)
        glossary[acronym.lower()] = definition
    for match in re.finditer(r'\b([A-Z]{2,})\s*[:\-–]\s*([A-Za-z ,&()/-]+)', full_text):
        acronym, definition = match.group(1), match.group(2).strip()
        glossary[acronym.lower()] = definition
    for match in re.finditer(r'\b([A-Z]{2,})\s+means\s+([A-Za-z ,&()/-]+)', full_text, re.IGNORECASE):
        acronym, definition = match.group(1), match.group(2).strip()
        glossary[acronym.lower()] = definition
    return glossary

def glossary_lookup(query, glossary=None):
    if glossary is None:
        glossary = get_global_glossary()
    q = query.lower().strip()
    patterns = [r"meaning of (\w+)", r"what is (\w+)", r"define (\w+)", r"full form of (\w+)", r"fullform of (\w+)"]
    for pat in patterns:
        m = re.search(pat, q)
        if m:
            term = m.group(1)
            if term in glossary:
                return f"{term.upper()} stands for '{glossary[term]}'."
    if q in glossary:
        return f"{q.upper()} stands for '{glossary[q]}'."
    return None

def save_glossary_to_cache(document_name, glossary):
    os.makedirs(GLOSSARY_CACHE_DIR, exist_ok=True)
    cache_path = os.path.join(GLOSSARY_CACHE_DIR, f"{document_name}_glossary.json")
    with open(cache_path, 'w', encoding='utf-8') as f:
        json.dump(glossary, f, indent=2, ensure_ascii=False)

def load_glossary_from_cache(document_name):
    cache_path = os.path.join(GLOSSARY_CACHE_DIR, f"{document_name}_glossary.json")
    if os.path.exists(cache_path):
        try:
            with open(cache_path, 'r', encoding='utf-8') as f:
                glossary = json.load(f)
            return glossary
        except Exception as e:
            print(f"[GLOSSARY] Error loading cache: {e}")
    return {}

def build_document_glossary(document_name):
    file_path = find_document_by_name(document_name)
    if not file_path:
        return {}
    cached_glossary = load_glossary_from_cache(document_name)
    if cached_glossary:
        return cached_glossary
    full_text = extract_text_from_file(file_path)
    if not full_text.strip():
        return {}
    glossary = extract_dynamic_glossary(full_text)
    save_glossary_to_cache(document_name, glossary)
    return glossary

def get_global_glossary():
    with _glossary_lock:
        if not _glossary_cache:
            all_docs = get_all_document_paths()
            combined_glossary = {}
            for doc in all_docs:
                doc_glossary = build_document_glossary(doc['filename'])
                combined_glossary.update(doc_glossary)
            _glossary_cache.update(combined_glossary)
        return _glossary_cache.copy()

def get_file_hash(file_path):
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def get_ocr_cache_path(file_path):
    os.makedirs(OCR_CACHE_DIR, exist_ok=True)
    file_hash = get_file_hash(file_path)
    filename = os.path.basename(file_path).replace('.pdf', '').replace('.docx', '').replace('.doc', '')
    cache_filename = f"{filename}_{file_hash}.txt"
    return os.path.join(OCR_CACHE_DIR, cache_filename)

def save_ocr_to_cache(file_path, ocr_text):
    cache_path = get_ocr_cache_path(file_path)
    metadata_path = cache_path.replace('.txt', '_metadata.json')
    with open(cache_path, 'w', encoding='utf-8') as f:
        f.write(ocr_text)
    metadata = {
        'original_file': file_path,
        'file_size': os.path.getsize(file_path),
        'ocr_date': datetime.now().isoformat(),
        'file_hash': get_file_hash(file_path),
        'text_length': len(ocr_text)
    }
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2)

def load_ocr_from_cache(file_path):
    cache_path = get_ocr_cache_path(file_path)
    metadata_path = cache_path.replace('.txt', '_metadata.json')
    if not os.path.exists(cache_path) or not os.path.exists(metadata_path):
        return None
    try:
        with open(metadata_path, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        current_hash = get_file_hash(file_path)
        if metadata.get('file_hash') != current_hash:
            return None
        with open(cache_path, 'r', encoding='utf-8') as f:
            ocr_text = f.read()
        return ocr_text
    except Exception as e:
        print(f"[CACHE] Error loading cache: {e}")
        return None

def extract_text_from_pdf_enhanced(file_path):
    cached_text = load_ocr_from_cache(file_path)
    if cached_text is not None:
        return cached_text
    try:
        doc = fitz.open(file_path)
        full_text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if len(text.strip()) < 100:
                blocks = page.get_text("blocks")
                text = "\n".join([block[4] for block in blocks if len(block) > 4])
            full_text += f"\n--- Page {page_num + 1} ---\n" + text
        doc.close()
        if len(full_text.strip()) > 200:
            save_ocr_to_cache(file_path, full_text)
        return full_text
    except Exception as e:
        print(f"[PDF] PyMuPDF failed: {e}")
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    try:
        images = convert_from_path(file_path, dpi=300)
        def process_page_enhanced(img_tuple):
            i, img = img_tuple
            custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
            text = pytesseract.image_to_string(img, lang='eng+hin', config=custom_config)
            return f"\n--- Page {i + 1} ---\n" + text
        with ThreadPoolExecutor(max_workers=3) as executor:
            results = list(executor.map(process_page_enhanced, enumerate(images)))
        full_text = "\n".join(results)
        save_ocr_to_cache(file_path, full_text)
        return full_text
    except Exception as e:
        print(f"[ERROR] Enhanced PDF extraction failed {file_path}: {e}")
        return ""

def extract_text_from_docx_with_cache(file_path):
    cached_text = load_ocr_from_cache(file_path)
    if cached_text is not None:
        return cached_text
    try:
        from docx import Document
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                full_text.append("\n--- TABLE ---\n" + "\n".join(table_text) + "\n--- END TABLE ---\n")
        text_result = "\n\n".join(full_text)
        save_ocr_to_cache(file_path, text_result)
        return text_result
    except Exception as e:
        print(f"[ERROR] Enhanced DOCX extraction failed {file_path}: {e}")
        return ""

def extract_text_from_txt(file_path):
    try:
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        print(f"[ERROR] Failed to extract from TXT {file_path}: {e}")
        return ""

def extract_text_from_image_with_cache(file_path):
    cached_text = load_ocr_from_cache(file_path)
    if cached_text is not None:
        return cached_text
    try:
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
        text = pytesseract.image_to_string(file_path, config=custom_config)
        save_ocr_to_cache(file_path, text)
        return text
    except Exception as e:
        print(f"[ERROR] Enhanced image extraction failed {file_path}: {e}")
        return ""

def extract_text_from_file(file_path):
    ocr_cache_path = get_ocr_cache_path(file_path)
    if os.path.exists(ocr_cache_path):
        with open(ocr_cache_path, 'r', encoding='utf-8') as f:
            text = f.read()
        print(f"[DEBUG] Loaded {len(text)} chars from OCR cache for {file_path}")
        return text
    # If not cached, extract and cache it
    if file_path.lower().endswith('.pdf'):
        text = extract_text_from_pdf_enhanced(file_path)
    elif file_path.lower().endswith('.docx') or file_path.lower().endswith('.doc'):
        text = extract_text_from_docx_with_cache(file_path)
    elif file_path.lower().endswith('.txt'):
        text = extract_text_from_txt(file_path)
    else:
        text = ""
    save_ocr_to_cache(file_path, text)
    print(f"[DEBUG] Extracted {len(text)} chars and cached for {file_path}")
    return text


def get_all_document_paths():
    document_paths = []
    base_paths = [
        "data/documents",
        "uploads",
        "user_uploads"
    ]
    for base_path in base_paths:
        if os.path.exists(base_path):
            for root, dirs, files in os.walk(base_path):
                for file in files:
                    if file.lower().endswith(('.pdf', '.docx', '.doc', '.txt')):
                        full_path = os.path.join(root, file)
                        rel_path = os.path.relpath(full_path, base_path)
                        # For data/documents, category is always the first folder (e.g. 'pgp' or 'hr')
                        if base_path == "data/documents":
                            parts = rel_path.split(os.sep)
                            category = parts[0] if len(parts) > 1 else ""
                        else:
                            category = ""
                        document_paths.append({
                            'filename': file,
                            'full_path': full_path,
                            'relative_path': rel_path,
                            'category': category
                        })
    return document_paths


def find_document_by_name(document_name):
    all_docs = get_all_document_paths()
    for doc in all_docs:
        if doc['filename'] == document_name:
            return doc['full_path']
    return None


# --- SEMANTIC SEARCH ACROSS CATEGORY ---

from shared_utils import *
def semantic_search_across_category(query, category, top_k=3):
    docs = get_all_document_paths()
    docs_in_cat = [doc for doc in docs if doc['category'].lower() == category.lower()]
    print(f"[SEMANTIC SEARCH] Searching {len(docs_in_cat)} docs in category '{category}' for query '{query}'")
    all_matches = []
    for doc in docs_in_cat:
        print(f"[SEMANTIC SEARCH] Checking doc: {doc['filename']}")
        matches = semantic_search(query, doc['filename'], top_k=top_k)
        print(f"[SEMANTIC SEARCH] Matches for {doc['filename']}: {matches}")
        for m in matches:
            all_matches.append((m, doc['filename']))
    all_matches = [x for x in all_matches if x[0].strip()]
    if not all_matches:
        print("[SEMANTIC SEARCH] No matches found.")
        return None
    best_chunk, best_doc = all_matches[0]
    print(f"[SEMANTIC SEARCH] Best match in {best_doc}: {best_chunk[:100]}")
    return f"From {best_doc}:\n\n{best_chunk.strip()}"

